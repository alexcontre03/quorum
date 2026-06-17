"""
build_memoria.py
================

Render the markdown chapters of a degree memoir into the corresponding Word
template (`MEMORIA_ING.docx` or `MEMORIA_CDIA.docx`), preserving the
template's front matter (cover, abstract, descriptors, index) and applying
the styles declared in the template.

Style mapping
-------------
- `# X. Title`   -> paragraph style ``Ttulo1`` (heading 1)
- `## X.Y ...`   -> paragraph style ``Ttulo2`` (heading 2)
- `### X.Y.Z ...`-> paragraph style ``Ttulo3`` (heading 3)
- prose paragraph -> ``Normal``
- bullet list (``- item``) -> ``Prrafodelista``
- markdown table -> native Word table

Inline runs
-----------
- ``*text*``  -> italic run
- `` `code` `` -> run in Consolas (monospaced)
- everything else -> plain run

Behaviour
---------
- Author notes (anything after ``---`` followed by ``### Author note ...``)
  are stripped per the project convention.
- Cross-references that reference a non-existent ``\\ref{...}`` are left as
  plain text (none expected in the .md sources).
- Everything in the .docx after the first ``Ttulo1`` paragraph is wiped
  before re-inserting the chapters, so the front matter (cover/abstract/
  index) stays intact while the body is rebuilt from scratch.

Run
---
    python scripts/build_memoria.py --grade ing
    python scripts/build_memoria.py --grade cdia
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
MEMORIA_ROOT = ROOT / "memoria"

# Per-degree configuration
CONFIG = {
    "ing": {
        "docx_path": MEMORIA_ROOT / "ingenieria" / "MEMORIA_ING.docx",
        "md_root": MEMORIA_ROOT / "ingenieria",
        "title": (
            "Design and implementation of a multi-agent software architecture "
            "for intelligent traceability of technical decisions in software development"
        ),
        "abstract_paragraphs": [
            "This project, within the Software Engineering field of study, designs and "
            "implements a multi-agent software architecture that enables intelligent "
            "traceability of technical decisions throughout the software development "
            "lifecycle. In collaborative development environments, important decisions "
            "taken during meetings or discussions are often disconnected from their later "
            "implementation in tasks, issues, or source code artifacts. The objective is "
            "to reduce that gap by creating a structured and scalable system capable of "
            "linking decisions to their corresponding development elements.",
            "The proposed solution is based on a modular architecture composed of "
            "specialised agents responsible for information processing, traceability "
            "management, and interaction with software development environments and "
            "related tools. From a software engineering perspective, the project focuses "
            "on system architecture design, backend implementation, inter-component "
            "communication, data persistence strategies, and integration through "
            "well-defined interfaces.",
            "Special attention is given to scalability, maintainability, robustness and "
            "extensibility of the architecture. The final outcome is a functional "
            "prototype demonstrating automated traceability workflows in realistic "
            "development scenarios, supported by technical validation and system-level "
            "testing. The work aligns with the competencies of Ingeniería Informática by "
            "focusing on software architecture design, distributed intelligent components, "
            "system integration and engineering best practices.",
        ],
        "descriptors": (
            "Multi-agent software architecture, Technical decision traceability, "
            "Local large language models, Closed-loop integration, Software engineering"
        ),
        "cover": {
            "degree_es": "Grado en Ingeniería Informática",
            "degree_eu": "Informatikako Ingeniaritzako gradua",
            "logo_image": MEMORIA_ROOT / "ingenieria" / "img" / "image1.jpeg",
            "banner_image": MEMORIA_ROOT / "ingenieria" / "img" / "image2.png",
        },
        "author": "Alejandro Contreras Alegria",
        "director": "Asier Perallos Ruiz",
        "place_date": "Bilbao, junio de 2026",
        "caption_position": "above",
        "chapter_files": [
            "01_introduction.md",
            "02_background.md",
            "03_objectives.md",
            "04_planning.md",
            "05_budget.md",
            "06_methodology.md",
            ("07_development", [
                "01_requirements.md",
                "02_design.md",
                "03_implementation.md",
                "04_testing.md",
                "05_closed_loop.md",
                "06_per_agent_models.md",
            ]),
            "08_ethics.md",
            "09_issues.md",
            "10_conclusions.md",
            "11_bibliography.md",
            "12_glossary.md",
            "13_appendices.md",
        ],
    },
    "cdia": {
        "docx_path": MEMORIA_ROOT / "cdia" / "MEMORIA_CDIA.docx",
        "md_root": MEMORIA_ROOT / "cdia",
        "title": (
            "Design and evaluation of generative LLM-based multi-agent systems for "
            "semantic reasoning and traceability in collaborative software engineering contexts"
        ),
        "abstract_paragraphs": [
            "This project, within the Generative AI field of study, designs and evaluates "
            "generative LLM-based multi-agent systems for semantic reasoning and "
            "traceability in collaborative software engineering contexts. The objective is "
            "to explore how generative language models can be orchestrated through "
            "coordinated agents to interpret unstructured textual information, identify "
            "relevant decisions and action items, and establish meaningful semantic links "
            "between related development artifacts.",
            "Special emphasis is placed on the design of agent coordination strategies "
            "and reasoning mechanisms enabled by large language models. The work involves "
            "experimentation with different generative configurations and coordination "
            "approaches in order to assess their effectiveness in complex traceability "
            "scenarios.",
            "A systematic evaluation framework is defined to measure performance in terms "
            "of semantic coherence, extraction accuracy and traceability precision. "
            "Controlled experimentation allows the analysis of different generative "
            "approaches and their impact on multi-agent collaboration and information "
            "quality.",
            "The final outcome includes an experimental validation of generative "
            "multi-agent architectures applied to decision traceability, together with a "
            "critical discussion of their strengths, limitations and potential "
            "improvements. The work directly aligns with the competencies of Ciencia de "
            "Datos e Inteligencia Artificial, particularly in generative AI, semantic "
            "modelling, multi-agent reasoning and applied experimentation with large "
            "language models.",
        ],
        "descriptors": (
            "Generative AI, Large language models, Multi-agent systems, "
            "Semantic traceability, Evaluation methodology"
        ),
        "cover": {
            "degree_es": "Grado en Ciencia de Datos e Inteligencia Artificial",
            "degree_eu": "Datuen Zientzia eta Adimen Artifizialeko gradua",
            "logo_image": MEMORIA_ROOT / "cdia" / "img" / "image1.jpeg",
            "banner_image": MEMORIA_ROOT / "cdia" / "img" / "image2.png",
        },
        "author": "Alejandro Contreras Alegria",
        "director": "Asier Perallos Ruiz",
        "place_date": "Bilbao, junio de 2026",
        "caption_position": "below",
        "chapter_files": [
            "01_introduction.md",
            "02_background.md",
            "03_objectives.md",
            "04_planning.md",
            "05_budget.md",
            "06_methodology.md",
            ("07_development", [
                "01_problem.md",
                "02_system_design.md",
                "03_rag.md",
                "04_dataset.md",
                "05_experimental.md",
                "06_results.md",
                "07_discussion.md",
            ]),
            "08_ethics.md",
            "09_issues.md",
            "10_conclusions.md",
            "11_bibliography.md",
            "12_glossary.md",
            "13_appendices.md",
        ],
    },
}

# Style names declared in styles.xml of the template.
STYLE_H1 = "heading 1"
STYLE_H2 = "heading 2"
STYLE_H3 = "heading 3"
STYLE_BODY = "Normal"
STYLE_LIST = "List Paragraph"

# Style ID for the heading 1 in this template, used to spot chapter starts.
STYLE_H1_ID = "Ttulo1"

CODE_FONT = "Consolas"


# ---------- inline parsing ----------------------------------------------

_INLINE_RE = re.compile(
    r"(?P<code>`[^`]+`)"
    r"|(?P<bold>\*\*[^*]+\*\*)"
    r"|(?P<italic>\*[^*]+\*)"
)


def add_runs(paragraph, text: str) -> None:
    """Append runs to *paragraph* honouring inline markdown (*, **, `)."""
    if not text:
        return
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group("code"):
            run = paragraph.add_run(m.group("code")[1:-1])
            run.font.name = CODE_FONT
        elif m.group("bold"):
            run = paragraph.add_run(m.group("bold")[2:-2])
            run.bold = True
        elif m.group("italic"):
            run = paragraph.add_run(m.group("italic")[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------- block parsing -----------------------------------------------

HEADING_RE = re.compile(r"^(#{1,4})\s+(.*?)\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|[\s\-:|]+\|\s*$")
FIGURE_RE = re.compile(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$")


def strip_heading_number(text: str) -> str:
    """Keep the leading numeric prefix like ``7.1.2`` in the heading title.

    The Deusto-style template renders ``1. INTRODUCTION`` and ``2.1 STATE OF
    THE ART``; the numeric prefix sits inside the title text, not as a
    separate auto-numbered field, so we preserve whatever the source .md
    declared.
    """
    return text.strip()


def parse_md(md_text: str) -> list[dict]:
    """Parse a markdown chapter file into a list of block descriptors.

    Each descriptor has a ``kind`` key (``h1``/``h2``/``h3``/``para``/``list``/
    ``table``) and the payload required to render it. Author-note blocks
    (anything after ``---`` followed by a ``### Author note`` heading) are
    stripped.
    """
    lines = md_text.splitlines()
    blocks: list[dict] = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # End of body: strip the trailing "Author note --- ..." section.
        if line.strip() == "---":
            # peek ahead: if the next non-empty line is an Author note heading,
            # stop processing.
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and lines[j].lstrip().startswith("### Author note"):
                break

        m = HEADING_RE.match(line)
        if m:
            hashes, title = m.group(1), strip_heading_number(m.group(2))
            level = len(hashes)
            if level == 1:
                blocks.append({"kind": "h1", "text": title})
            elif level == 2:
                blocks.append({"kind": "h2", "text": title})
            elif level == 3:
                blocks.append({"kind": "h3", "text": title})
            else:
                blocks.append({"kind": "h3", "text": title})
            i += 1
            continue

        # figure: ![caption](path)
        m_fig = FIGURE_RE.match(line)
        if m_fig:
            blocks.append({"kind": "figure", "caption": m_fig.group(1), "path": m_fig.group(2)})
            i += 1
            continue

        # markdown table?
        if line.lstrip().startswith("|") and i + 1 < len(lines) and TABLE_SEP_RE.match(lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip header + separator
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            blocks.append({"kind": "table", "header": header, "rows": rows})
            continue

        # bullet list
        if line.lstrip().startswith("- "):
            items = []
            while i < len(lines) and lines[i].lstrip().startswith("- "):
                items.append(lines[i].lstrip()[2:].strip())
                i += 1
            blocks.append({"kind": "list", "items": items})
            continue

        # blank line: skip
        if not line.strip():
            i += 1
            continue

        # paragraph: gather until blank line or new block
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not HEADING_RE.match(lines[i]):
            if lines[i].strip() == "---":
                break
            if lines[i].lstrip().startswith("- "):
                break
            if lines[i].lstrip().startswith("|"):
                break
            para_lines.append(lines[i])
            i += 1
        blocks.append({"kind": "para", "text": " ".join(s.strip() for s in para_lines)})

    return blocks


# ---------- document mutation -------------------------------------------

def _build_cover_paragraphs(doc, cfg: dict, signed: bool = False) -> list:
    """Build one Deusto cover page as an ordered list of OxmlElement paragraphs.

    The university normative mandates two identical cover pages: an external
    one (``signed=False``) and an internal one (``signed=True``) that holds
    a signature line for the director. The bilingual headings stay the same
    in both; only the signing block differs.
    """
    cov = cfg["cover"]
    paragraphs: list = []

    def make_p(text: str = "", *, size_pt: int | None = None,
               bold: bool = False, italic: bool = False, align: str = "center",
               space_before_pt: int | None = None,
               space_after_pt: int | None = None,
               image_path: Path | None = None, image_width_cm: float | None = None):
        p = doc.add_paragraph()
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if space_before_pt is not None:
            p.paragraph_format.space_before = Pt(space_before_pt)
        if space_after_pt is not None:
            p.paragraph_format.space_after = Pt(space_after_pt)
        if image_path is not None and image_path.exists():
            run = p.add_run()
            if image_width_cm is not None:
                run.add_picture(str(image_path), width=Cm(image_width_cm))
            else:
                run.add_picture(str(image_path))
        elif text:
            run = p.add_run(text)
            if size_pt is not None:
                run.font.size = Pt(size_pt)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        # Detach from the body (we will re-attach in order at the top).
        body = doc.element.body
        body.remove(p._p)
        paragraphs.append(p._p)

    # Block 1: Deusto logo
    make_p(image_path=cov["logo_image"], image_width_cm=5.5, space_after_pt=18)

    # Block 2: degree name (bilingual)
    make_p(cov["degree_es"], size_pt=16, bold=True, space_after_pt=2)
    make_p(cov["degree_eu"], size_pt=14, space_after_pt=24)

    # Block 3: PFG banner
    make_p(image_path=cov["banner_image"], image_width_cm=10, space_after_pt=6)

    # Block 4: PFG label (bilingual)
    make_p("Proyecto fin de grado", size_pt=14, bold=True, space_after_pt=2)
    make_p("Gradu amaierako proiektua", size_pt=12, space_after_pt=80)

    # Block 5: project title
    make_p(cfg["title"], size_pt=16, bold=True, space_after_pt=60)

    # Block 6: author / director / place-date
    make_p(cfg["author"], size_pt=12, space_after_pt=6)
    make_p(f"Director: {cfg['director']}", size_pt=12, space_after_pt=6)
    make_p(cfg["place_date"], size_pt=12, space_after_pt=0)

    # Signature block (internal cover only)
    if signed:
        make_p("", size_pt=10, space_before_pt=36, space_after_pt=2)
        make_p("____________________________", size_pt=12, space_after_pt=2)
        make_p("Firma del Director / Zuzendariaren sinadura",
               size_pt=10, italic=True, space_after_pt=0)

    # Final page break
    pb_p = OxmlElement("w:p")
    pb_r = OxmlElement("w:r")
    pb_br = OxmlElement("w:br")
    pb_br.set(qn("w:type"), "page")
    pb_r.append(pb_br)
    pb_p.append(pb_r)
    paragraphs.append(pb_p)

    return paragraphs


def insert_cover_page(doc, cfg: dict) -> None:
    """Insert a native Word cover page at the top of *doc*.

    Idempotent: looks for the bookmark ``cover_inserted`` and skips if found.
    Removes any previous cover (paragraphs before the first ``Subttulo``) so
    re-running the build is safe.
    """
    body = doc.element.body

    for bm in body.iter(qn("w:bookmarkStart")):
        if bm.get(qn("w:name")) == "cover_inserted":
            # Already present: remove everything before the first Subttulo,
            # then re-insert. That way edits to the cover config propagate.
            children = list(body.iterchildren())
            for ch in children:
                if ch.tag != qn("w:p"):
                    continue
                pstyle = ch.find(qn("w:pPr") + "/" + qn("w:pStyle"))
                sid = pstyle.get(qn("w:val")) if pstyle is not None else None
                if sid == "Subttulo":
                    break
                body.remove(ch)
            break

    cover_paragraphs = (
        _build_cover_paragraphs(doc, cfg, signed=False)
        + _build_cover_paragraphs(doc, cfg, signed=True)
    )

    # Insert in order at the very top of the body.
    anchor = body[0] if len(body) else None

    # Wrap the first paragraph with the bookmark for idempotency.
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), "9000")
    bm_start.set(qn("w:name"), "cover_inserted")
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), "9000")
    cover_paragraphs[0].insert(0, bm_start)
    cover_paragraphs[0].append(bm_end)

    for p in cover_paragraphs:
        if anchor is not None:
            anchor.addprevious(p)
        else:
            body.append(p)


def _insert_field(paragraph, instr: str, default_text: str = "Update field (F9)") -> None:
    """Insert a Word complex field with *instr* as its instruction text.

    The field is rendered with a placeholder body so that Word displays
    something readable when fields have not been refreshed yet. Word
    recomputes the field on open when ``updateFields`` is ``true`` in
    settings.xml.
    """
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r.append(begin)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r.append(it)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r.append(sep)
    t = OxmlElement("w:t")
    t.text = default_text
    r.append(t)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(end)


def _make_subtitle(doc, label: str):
    """Create and return a new paragraph with the ``Subttulo`` style."""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles["Subtitle"]
    except KeyError:
        pass
    p.add_run(label)
    body = doc.element.body
    body.remove(p._p)
    return p._p


def _make_field_paragraph(doc, instr: str, default_text: str):
    p = doc.add_paragraph()
    _insert_field(p, instr, default_text)
    body = doc.element.body
    body.remove(p._p)
    return p._p


def insert_toc_lof_lot(doc) -> None:
    """Replace the Index placeholder with a TOC field and add LOF + LOT.

    The Deusto template ships with a ``Subttulo`` paragraph reading "Índice"
    followed by two informative ``Normal`` placeholders. We:
    - keep the Subtitle (already renamed to "Index"),
    - drop the informative placeholders,
    - insert a TOC field that Word fills on open,
    - then add new Subtitles "List of figures" and "List of tables" and the
      corresponding TOC field for each one.
    """
    body = doc.element.body
    children = list(body.iterchildren())

    # Find the Index subtitle
    idx_index = None
    for i, ch in enumerate(children):
        if ch.tag != qn("w:p"):
            continue
        pstyle = ch.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        if pstyle is not None and pstyle.get(qn("w:val")) == "Subttulo":
            text = "".join(t.text or "" for t in ch.iter(qn("w:t"))).strip().lower()
            if text in ("index", "índice", "indice"):
                idx_index = i
                break
    if idx_index is None:
        return

    # Remove the Normal placeholders that follow Index (until the first Ttulo1
    # or Subttulo). Any sectPr that the template embedded in those
    # placeholders has already been promoted to a body child by
    # ``ensure_final_sectpr``, so the deletion is safe.
    cur = idx_index + 1
    while cur < len(children):
        nxt = children[cur]
        if nxt.tag != qn("w:p"):
            break
        pstyle = nxt.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        sid = pstyle.get(qn("w:val")) if pstyle is not None else None
        if sid in ("Ttulo1", "Subttulo"):
            break
        body.remove(nxt)
        cur += 1

    anchor = children[idx_index]

    # Re-fetch the next sibling because the children list is stale after removals.
    def add_after(anchor_elem, new_elem):
        anchor_elem.addnext(new_elem)
        return new_elem

    toc_p = _make_field_paragraph(doc, 'TOC \\o "1-3" \\h \\z \\u', "Right-click and choose 'Update Field' to populate the table of contents.")
    last = add_after(anchor, toc_p)

    lof_sub = _make_subtitle(doc, "List of figures")
    last = add_after(last, lof_sub)
    lof_field = _make_field_paragraph(doc, 'TOC \\h \\z \\c "Figure"', "Right-click and choose 'Update Field' to populate the list of figures.")
    last = add_after(last, lof_field)

    lot_sub = _make_subtitle(doc, "List of tables")
    last = add_after(last, lot_sub)
    lot_field = _make_field_paragraph(doc, 'TOC \\h \\z \\c "Table"', "Right-click and choose 'Update Field' to populate the list of tables.")
    add_after(last, lot_field)


def _build_sect_pr(*, fmt: str | None = None, start: int | None = None,
                   ref_settings: dict | None = None,
                   page_size: tuple[str, str] = ("11907", "16840"),
                   margins: dict | None = None,
                   header_ref_id: str | None = None,
                   footer_ref_id: str | None = None):
    """Build a w:sectPr OxmlElement with the given pagination settings."""
    sectPr = OxmlElement("w:sectPr")

    if header_ref_id is not None:
        hr = OxmlElement("w:headerReference")
        hr.set(qn("w:type"), "default")
        hr.set(qn("r:id"), header_ref_id)
        sectPr.append(hr)

    if footer_ref_id is not None:
        fr = OxmlElement("w:footerReference")
        fr.set(qn("w:type"), "default")
        fr.set(qn("r:id"), footer_ref_id)
        sectPr.append(fr)

    pg_sz = OxmlElement("w:pgSz")
    pg_sz.set(qn("w:w"), page_size[0])
    pg_sz.set(qn("w:h"), page_size[1])
    sectPr.append(pg_sz)

    pg_mar = OxmlElement("w:pgMar")
    m = margins or {"top": "1985", "right": "1418", "bottom": "1418",
                    "left": "1418", "header": "567", "footer": "283", "gutter": "0"}
    for k, v in m.items():
        pg_mar.set(qn(f"w:{k}"), v)
    sectPr.append(pg_mar)

    if fmt or start is not None:
        pg_num = OxmlElement("w:pgNumType")
        if fmt is not None:
            pg_num.set(qn("w:fmt"), fmt)
        if start is not None:
            pg_num.set(qn("w:start"), str(start))
        sectPr.append(pg_num)

    cols = OxmlElement("w:cols")
    cols.set(qn("w:space"), "720")
    sectPr.append(cols)

    return sectPr


def setup_pagination(doc) -> None:
    """Configure three sections in the document:

    1. Covers (no page numbers).
    2. Front matter (Abstract → List of tables): lower-roman, start=1.
    3. Body (Chapter 1 onwards): decimal, start=1.

    The post-processing on the ZIP that follows this function adds a blank
    footer for section 1 so that pages 1 and 2 do not display a number.
    """
    body = doc.element.body

    # Locate the "Abstract" Subtitle (first paragraph after the covers) and
    # the first Heading 1 (first chapter).
    children = list(body.iterchildren())
    idx_abstract = None
    idx_chapter = None
    for i, ch in enumerate(children):
        if ch.tag != qn("w:p"):
            continue
        pstyle = ch.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        if pstyle is None:
            continue
        sid = pstyle.get(qn("w:val"))
        if sid == "Subttulo" and idx_abstract is None:
            text = "".join(t.text or "" for t in ch.iter(qn("w:t"))).strip().lower()
            if text in ("abstract", "resumen"):
                idx_abstract = i
        if sid == "Ttulo1" and idx_chapter is None:
            idx_chapter = i
            break

    # Section 1 ends in the paragraph just before idx_abstract.
    # Section 2 ends in the paragraph just before idx_chapter.
    # We attach a sectPr to those "boundary" paragraphs by putting it inside
    # their w:pPr element.

    def attach_sectpr(target_p, sectPr):
        ppr = target_p.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            target_p.insert(0, ppr)
        # remove existing sectPr if any
        existing = ppr.find(qn("w:sectPr"))
        if existing is not None:
            ppr.remove(existing)
        ppr.append(sectPr)

    if idx_abstract is not None and idx_abstract > 0:
        # The page-break paragraph that closes the internal cover sits just
        # before idx_abstract. We attach the cover sectPr there.
        cover_close = children[idx_abstract - 1]
        cover_sect = _build_sect_pr()  # no numFmt = no number on covers
        attach_sectpr(cover_close, cover_sect)

    if idx_chapter is not None and idx_chapter > 0:
        front_close = children[idx_chapter - 1]
        front_sect = _build_sect_pr(fmt="lowerRoman", start=1)
        attach_sectpr(front_close, front_sect)

    # Update the final sectPr (last direct child of body, after the chapter
    # bodies) to decimal start=1.
    final_sect = None
    for child in body.iterchildren():
        if child.tag == qn("w:sectPr"):
            final_sect = child
    if final_sect is not None:
        pgnum = final_sect.find(qn("w:pgNumType"))
        if pgnum is None:
            pgnum = OxmlElement("w:pgNumType")
            final_sect.append(pgnum)
        pgnum.set(qn("w:fmt"), "decimal")
        pgnum.set(qn("w:start"), "1")
    else:
        print("  WARNING: no final sectPr found; body pagination not configured")


def export_to_pdf(docx_path: Path) -> Path:
    """Open *docx_path* in Word, force a full field update (TOC, LOF, LOT,
    SEQ) and export to a PDF next to the .docx.

    Requires Microsoft Word installed and ``pywin32`` available. Word must
    not have the file already open.
    """
    import win32com.client as win32  # noqa: F401  (delayed import)

    pdf_path = docx_path.with_suffix(".pdf")
    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path.resolve()))
        # First pass: settle SEQ values across the body.
        doc.Fields.Update()
        # Second pass: rebuild TOC, LOF and LOT against the now-correct SEQs.
        for toc in doc.TablesOfContents:
            toc.Update()
        for tof in doc.TablesOfFigures:
            tof.Update()
        # Third pass: catch any remaining field that depended on the second one.
        doc.Fields.Update()
        for toc in doc.TablesOfContents:
            toc.Update()
        for tof in doc.TablesOfFigures:
            tof.Update()
        # 17 = wdExportFormatPDF
        doc.SaveAs2(str(pdf_path.resolve()), FileFormat=17)
        doc.Close(SaveChanges=False)
        return pdf_path
    finally:
        word.Quit()


def finalize_docx_indices(docx_path: Path) -> None:
    """Open *docx_path* in Word, force a full field update on all TOCs and
    PAGEREFs, save the docx, close.

    This is what makes the manually-built indices (TOC, LOF, LOT) carry the
    actual page numbers instead of the placeholder "1" that Word shows for
    fields that have not been refreshed. Without this pass, the user has to
    press Ctrl+A then F9 inside Word and confirm the update of each table;
    even then, Word does not always recompute correctly because the Deusto
    template ships heading styles (Ttulo1, Ttulo2, Ttulo3) whose outline
    levels Word does not always read on F9.

    Requires Microsoft Word installed and ``pywin32`` available. Word must
    not have the file already open.
    """
    import win32com.client as win32  # noqa: F401

    # Late-binding Dispatch avoids the gen_py cache, which is prone to
    # corruption ("CLSIDToClassMap" attribute errors) when the Word version
    # changes underneath an installed pywin32.
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path.resolve()))
        # Force Word to paginate the document before reading any page number.
        # Without this, every Field.Update() reads a stale page count and the
        # TOC entries collapse onto the page where the body section starts
        # (typically page 3, which is the canonical "all entries say 3"
        # symptom). Repaginate() flushes the layout engine.
        doc.Repaginate()
        # Three-pass update mirrors the PDF export discipline: first the
        # leaf SEQ fields, then the tables that depend on them, then any
        # PAGEREF that depends on the now-correct TOC pagination.
        doc.Fields.Update()
        for i in range(1, doc.TablesOfContents.Count + 1):
            doc.TablesOfContents(i).Update()
        for i in range(1, doc.TablesOfFigures.Count + 1):
            doc.TablesOfFigures(i).Update()
        doc.Repaginate()
        doc.Fields.Update()
        for i in range(1, doc.TablesOfContents.Count + 1):
            doc.TablesOfContents(i).Update()
        for i in range(1, doc.TablesOfFigures.Count + 1):
            doc.TablesOfFigures(i).Update()
        # Save the docx in place so the user opens an already-populated copy.
        doc.Save()
        doc.Close(SaveChanges=False)
    finally:
        word.Quit()


def enable_update_fields(docx_path: Path) -> None:
    """Patch settings.xml inside *docx_path* to add ``w:updateFields=true``
    and strip the ``numPr`` (auto-numbering) from Heading 1/2/3 styles so
    Word does not duplicate the chapter numbers that the source markdown
    already provides in the heading text.
    """
    import zipfile
    import shutil

    tmp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == "word/settings.xml":
                s = data.decode("utf-8")
                if "updateFields" not in s:
                    s = re.sub(
                        r"(<w:settings[^>]*>)",
                        r'\1<w:updateFields w:val="true"/>',
                        s,
                        count=1,
                    )
                data = s.encode("utf-8")
            elif item == "word/styles.xml":
                s = data.decode("utf-8")
                # Remove <w:numPr>...</w:numPr> blocks from Ttulo1, Ttulo2 and
                # Ttulo3 style definitions only.
                for style_id in ("Ttulo1", "Ttulo2", "Ttulo3"):
                    pattern = rf'(<w:style[^>]*w:styleId="{style_id}"[^>]*>.*?)<w:numPr>.*?</w:numPr>(.*?</w:style>)'
                    s = re.sub(pattern, r"\1\2", s, flags=re.DOTALL)
                data = s.encode("utf-8")
            zout.writestr(item, data)
    shutil.move(str(tmp_path), str(docx_path))


def _rename_subtitle(doc, old_label: str, new_label: str) -> None:
    """Rename a Subtitle paragraph whose visible text equals *old_label*."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        pstyle = child.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        if pstyle is None or pstyle.get(qn("w:val")) != "Subttulo":
            continue
        text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
        if text.lower() != old_label.lower():
            continue
        ts = list(child.iter(qn("w:t")))
        if not ts:
            continue
        ts[0].text = new_label
        for extra in ts[1:]:
            extra.text = ""
        return


def ensure_final_sectpr(doc) -> None:
    """Guarantee that the body has exactly one ``w:sectPr`` as its last child.

    The Deusto template ships with the template's final ``sectPr`` buried
    inside one of the ``Index`` placeholder paragraphs' ``w:pPr``. Several
    python-docx helpers (e.g. ``Document.add_table``) crash when the body
    has no direct sectPr, because they read ``Document.sections[-1]`` to
    compute the block width. We promote any embedded sectPr to a direct
    child of the body at the very end so subsequent helpers find what they
    expect.
    """
    body = doc.element.body
    embedded = None
    for p in body.iter(qn("w:p")):
        ppr = p.find(qn("w:pPr"))
        if ppr is None:
            continue
        sect = ppr.find(qn("w:sectPr"))
        if sect is not None:
            embedded = sect
            ppr.remove(sect)
            break

    direct = None
    for child in body.iterchildren():
        if child.tag == qn("w:sectPr"):
            direct = child
            break

    if direct is None:
        if embedded is not None:
            body.append(embedded)
        else:
            # Fall back to a minimal default sectPr; the page size and margins
            # mirror the Deusto template defaults.
            body.append(_build_sect_pr())


def replace_front_matter(doc, abstract_paragraphs: list[str], descriptors: str) -> None:
    """Replace the body paragraphs of the ``Resumen`` and ``Descriptores`` sections.

    The Deusto template ships a placeholder under each ``Subtitle`` heading
    (``Resumen``, ``Descriptores``, ``Índice``). We keep the subtitles
    untouched and overwrite the paragraphs sitting between one subtitle and
    the next with the real abstract / descriptors content.
    """
    body = doc.element.body
    children = list(body.iterchildren())

    def find_subtitle_run(label_candidates: list[str]) -> int | None:
        for idx, child in enumerate(children):
            if child.tag != qn("w:p"):
                continue
            pstyle = child.find(qn("w:pPr") + "/" + qn("w:pStyle"))
            style_id = pstyle.get(qn("w:val")) if pstyle is not None else None
            text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
            if style_id == "Subttulo" and any(text.lower().startswith(c.lower()) for c in label_candidates):
                return idx
        return None

    def collect_following_normals(start_idx: int) -> list:
        """Return the contiguous Normal/Default paragraphs after start_idx."""
        out = []
        i = start_idx + 1
        while i < len(children):
            child = children[i]
            if child.tag != qn("w:p"):
                break
            pstyle = child.find(qn("w:pPr") + "/" + qn("w:pStyle"))
            sid = pstyle.get(qn("w:val")) if pstyle is not None else None
            if sid == "Subttulo" or sid == "Ttulo1":
                break
            out.append(child)
            i += 1
        return out

    # Rename Spanish front-matter labels to English since the body is in
    # English (the bilingual covers stay in ES/EU as the normative requires).
    _rename_subtitle(doc, "Resumen", "Abstract")
    _rename_subtitle(doc, "Descriptores", "Keywords")
    _rename_subtitle(doc, "Índice", "Index")
    # Refresh the children list after the renames.
    children = list(body.iterchildren())

    # Replace Abstract content
    idx_resumen = find_subtitle_run(["Abstract", "Resumen"])
    if idx_resumen is not None:
        placeholders = collect_following_normals(idx_resumen)
        anchor = children[idx_resumen]
        # remove the existing placeholders
        for ph in placeholders:
            body.remove(ph)
        # insert new paragraphs right after the subtitle
        from docx.oxml import OxmlElement

        cursor = anchor
        for txt in abstract_paragraphs:
            new_p = OxmlElement("w:p")
            anchor.addnext(new_p) if cursor is anchor else cursor.addnext(new_p)
            # Use the standard add_paragraph workflow via a temporary doc trick:
            # build runs manually.
            ppr = OxmlElement("w:pPr")
            pstyle = OxmlElement("w:pStyle")
            pstyle.set(qn("w:val"), "Normal")
            ppr.append(pstyle)
            new_p.append(ppr)
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = txt
            t.set(qn("xml:space"), "preserve")
            r.append(t)
            new_p.append(r)
            cursor = new_p
        # refresh children for the next lookup
        children = list(body.iterchildren())

    # Replace Keywords
    idx_desc = find_subtitle_run(["Keywords", "Descriptores"])
    if idx_desc is not None:
        placeholders = collect_following_normals(idx_desc)
        anchor = children[idx_desc]
        for ph in placeholders:
            body.remove(ph)
        from docx.oxml import OxmlElement

        new_p = OxmlElement("w:p")
        anchor.addnext(new_p)
        ppr = OxmlElement("w:pPr")
        pstyle = OxmlElement("w:pStyle")
        pstyle.set(qn("w:val"), "Normal")
        ppr.append(pstyle)
        new_p.append(ppr)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = descriptors
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        new_p.append(r)


def wipe_body_completely(doc) -> None:
    """Delete every paragraph and table from the body, preserving only the
    final ``w:sectPr`` element that holds page size and margins.

    The previous helper preserved everything before the first chapter
    heading, but the Deusto template ships with two cover pages and the
    Resumen/Descriptores/Indice placeholders embedded in the body. Repeated
    rebuilds duplicated those placeholders and never removed the covers.
    This helper starts from a clean slate so the front matter we build later
    is exactly what we asked for and nothing else.
    """
    body = doc.element.body
    # Hunt down the final sectPr (either a direct child of body or embedded
    # inside the last paragraph's pPr). Preserve exactly one.
    preserved_sectpr = None
    for ch in list(body.iterchildren()):
        if ch.tag == qn("w:sectPr"):
            preserved_sectpr = ch
            break
    if preserved_sectpr is None:
        # Look inside paragraphs.
        for p in body.iter(qn("w:p")):
            ppr = p.find(qn("w:pPr"))
            if ppr is not None:
                sect = ppr.find(qn("w:sectPr"))
                if sect is not None:
                    preserved_sectpr = sect
                    ppr.remove(sect)
                    break
    # Remove every direct child of body.
    for ch in list(body.iterchildren()):
        body.remove(ch)
    # Reinsert the sectPr at the end. If we found none, fall back to a default.
    if preserved_sectpr is not None:
        body.append(preserved_sectpr)
    else:
        body.append(_build_sect_pr())


def add_paragraph(doc, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    # Force every top-level chapter to start on a fresh page, with the
    # exception of the very first one (which already follows the LOT). Without
    # this, a chapter that ends near the page bottom runs straight into the
    # next chapter's title, which the academic format does not allow.
    if style == STYLE_H1:
        H1_COUNTER[0] += 1
        if H1_COUNTER[0] > 1:
            p.paragraph_format.page_break_before = True
    add_runs(p, text)

    # Drop a bookmark on every heading so the manually built index can point
    # at it via PAGEREF, and stamp the explicit outline level so Word's TOC
    # field (TOC \u) discovers the heading regardless of whether the heading
    # style itself (Ttulo1/Ttulo2/Ttulo3 in the Deusto template) carries the
    # outline level on its definition. Without this, the TOC field would
    # build entries with placeholder page numbers because Word would not
    # find any paragraph at outline level 1, 2 or 3.
    if style in (STYLE_H1, STYLE_H2, STYLE_H3):
        level = {STYLE_H1: 1, STYLE_H2: 2, STYLE_H3: 3}[style]
        bm = _new_bookmark_name(f"h{level}")
        _wrap_paragraph_with_bookmark(p._p, bm)
        TOC_ENTRIES.append((level, text, bm))
        # Set outline level on the paragraph (0-based: H1 -> "0").
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._p.insert(0, pPr)
        outline = pPr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            pPr.append(outline)
        outline.set(qn("w:val"), str(level - 1))


def _add_table_borders(table) -> None:
    """Add a thin black border around every cell of *table*."""
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)


def add_table(doc, header: list[str], rows: list[list[str]]) -> None:
    if not header:
        return
    TABLE_COUNTER[0] += 1
    bm_name = _new_bookmark_name(f"tbl_{TABLE_COUNTER[0]}")
    tbl_number = TABLE_COUNTER[0]
    # Derive a short summary for the LOT from the header row.
    summary = " / ".join(h for h in header if h)
    LOT_ENTRIES.append((f"Table {tbl_number}: {summary}", bm_name))
    # caption paragraph for the table. We use Normal style with manual
    # formatting (same trick as figure captions) so a long summary is not
    # truncated by the template's "Figura" style.
    cap_p = doc.add_paragraph()
    try:
        cap_p.style = doc.styles["Normal"]
    except KeyError:
        pass
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(12)
    cap_p.paragraph_format.space_after = Pt(6)
    cap_p.paragraph_format.keep_with_next = True
    cap_p.paragraph_format.keep_together = False
    label_run = cap_p.add_run(f"Table {tbl_number}: ")
    label_run.italic = True
    label_run.bold = True
    body_run = cap_p.add_run(_plain_text_for_caption(summary))
    body_run.italic = True
    _wrap_paragraph_with_bookmark(cap_p._p, bm_name)

    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    # The Deusto template does not define ``Table Grid``. We add borders
    # manually so the rendered table is not invisible.
    _add_table_borders(table)
    for j, cell_text in enumerate(header):
        cell = table.cell(0, j)
        cell.text = ""
        add_runs(cell.paragraphs[0], cell_text)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows, start=1):
        # tolerate ragged rows
        for j in range(len(header)):
            cell = table.cell(i, j)
            cell.text = ""
            content = row[j] if j < len(row) else ""
            add_runs(cell.paragraphs[0], content)
    # add a spacer paragraph after the table so consecutive blocks breathe and
    # the next paragraph does not start glued to the last row.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(12)


FIGURE_COUNTER = [0]
TABLE_COUNTER = [0]
H1_COUNTER = [0]
BOOKMARK_COUNTER = [1000]
FIGURE_STYLE = "Figura"
CAPTION_POSITION = ["below"]

# Lists captured during body rendering, consumed when building the front
# matter index, figure index and table index.
TOC_ENTRIES: list[tuple[int, str, str]] = []   # (level, title, bookmark)
LOF_ENTRIES: list[tuple[str, str]] = []        # (caption, bookmark)
LOT_ENTRIES: list[tuple[str, str]] = []        # (label,   bookmark)


def _new_bookmark_name(prefix: str) -> str:
    BOOKMARK_COUNTER[0] += 1
    return f"bm_{prefix}_{BOOKMARK_COUNTER[0]}"


def _wrap_paragraph_with_bookmark(p_element, bm_name: str) -> None:
    """Insert ``w:bookmarkStart`` and ``w:bookmarkEnd`` around the runs of the
    paragraph so a Word PAGEREF field can later jump to it."""
    bm_id = str(BOOKMARK_COUNTER[0])
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bm_id)
    start.set(qn("w:name"), bm_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bm_id)
    # Insert start right after the pPr (if any) and end at the end.
    ppr = p_element.find(qn("w:pPr"))
    if ppr is not None:
        ppr.addnext(start)
    else:
        p_element.insert(0, start)
    p_element.append(end)


def _insert_pageref_field(paragraph, bookmark_name: str) -> None:
    """Append a PAGEREF field run that resolves to the page of *bookmark_name*."""
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' PAGEREF {bookmark_name} \\h '
    r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r.append(sep)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(end)


def _add_seq_field(paragraph, seq_name: str) -> None:
    """Append a Word SEQ field run to *paragraph*.

    SEQ fields are what Word uses to auto-number figures and tables; the
    same fields are picked up by the ``TOC \\c "Figure"`` / ``TOC \\c "Table"``
    list fields when generating the list of figures / tables.
    """
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' SEQ {seq_name} \\* ARABIC '
    r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r.append(sep)
    t = OxmlElement("w:t")
    t.text = "1"  # placeholder; Word recomputes on open
    r.append(t)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r.append(end)


def _plain_text_for_caption(caption: str) -> str:
    """Strip inline markdown (backticks, asterisks) from the caption so the
    body of the figure/table label renders as a single, italicised string.
    Captions in the source markdown sometimes wrap technical identifiers in
    backticks (`task_followup_agent`); the backticks themselves are noise in
    the rendered caption and would otherwise force a font switch in the
    middle of the paragraph."""
    out = re.sub(r"`([^`]+)`", r"\1", caption)
    out = re.sub(r"\*\*([^*]+)\*\*", r"\1", out)
    out = re.sub(r"\*([^*]+)\*", r"\1", out)
    return out


def _figure_dimensions(img_path: Path) -> tuple[float, float]:
    """Return ``(width_cm, height_cm)`` for an image, capping the height at
    18 cm and the width at 14 cm. This avoids the failure mode where a tall
    image scaled to 14 cm wide grew past one page and overlapped the
    surrounding text."""
    try:
        from PIL import Image as PILImage
        with PILImage.open(img_path) as im:
            w_px, h_px = im.size
    except Exception:
        return (14.0, 8.0)
    if w_px <= 0 or h_px <= 0:
        return (14.0, 8.0)
    MAX_W_CM = 14.0
    MAX_H_CM = 18.0
    ratio = h_px / w_px
    target_w = MAX_W_CM
    target_h = target_w * ratio
    if target_h > MAX_H_CM:
        target_h = MAX_H_CM
        target_w = target_h / ratio
    return (target_w, target_h)


def add_figure(doc, caption: str, path_str: str, md_root: Path) -> None:
    img_path = md_root / path_str if not Path(path_str).is_absolute() else Path(path_str)
    FIGURE_COUNTER[0] += 1
    bm_name = _new_bookmark_name(f"fig_{FIGURE_COUNTER[0]}")
    fig_number = FIGURE_COUNTER[0]
    # Strip backtick code spans from the caption used in the index so the LOF
    # reads as plain text instead of carrying Markdown code markers.
    plain_caption = re.sub(r"`([^`]+)`", r"\1", caption)
    LOF_ENTRIES.append((f"Figure {fig_number}: {plain_caption}", bm_name))

    def make_caption(*, keep_with_next: bool, bookmark: str | None = None):
        # Use the body Normal style instead of the template's "Figura" style.
        # The "Figura" style ships with ``keepLines`` enabled and forces every
        # line of the paragraph to stay on the same page, which Word silently
        # truncates when the caption is long. Plain Normal + italic + centred
        # renders the full caption every time.
        p = doc.add_paragraph()
        try:
            p.style = doc.styles["Normal"]
        except KeyError:
            pass
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.keep_with_next = keep_with_next
        # ``keep_together`` is intentionally False on captions so Word does not
        # truncate them when they are long; the figure above (or below) carries
        # the keep_together flag instead.
        p.paragraph_format.keep_together = False
        label_run = p.add_run(f"Figure {fig_number}: ")
        label_run.italic = True
        label_run.bold = True
        # Body of the caption rendered in italic.
        body_run = p.add_run(_plain_text_for_caption(caption))
        body_run.italic = True
        if bookmark is not None:
            _wrap_paragraph_with_bookmark(p._p, bookmark)
        return p

    def make_image(*, keep_with_next: bool, bookmark: str | None = None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        # Force the line to grow with the image instead of staying at the
        # paragraph style's fixed line height. Without this, an inline image
        # taller than one line gets clipped to a narrow band and the next
        # paragraph starts inside the image's vertical extent.
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # ``keep_together`` keeps the image inside one page boundary, but we
        # do not want ``keep_with_next`` on tall images (it forces an awkward
        # page break above the image when the caption is below). The caller
        # passes the right value for the position of the caption.
        p.paragraph_format.keep_with_next = keep_with_next
        p.paragraph_format.keep_together = True
        # Resolve dimensions that fit on the page. Wide-and-flat figures end
        # up narrow in height; tall figures are capped at 18 cm of height so
        # they fit on a single page.
        if img_path.exists():
            w_cm, h_cm = _figure_dimensions(img_path)
            p.add_run().add_picture(str(img_path), width=Cm(w_cm), height=Cm(h_cm))
        else:
            p.add_run(f"[Missing image: {path_str}]").italic = True
        if bookmark is not None:
            _wrap_paragraph_with_bookmark(p._p, bookmark)
        return p

    # The bookmark is attached to whichever element comes first (caption above
    # vs image below), so PAGEREF resolves to the topmost page of the figure.
    if CAPTION_POSITION[0] == "above":
        make_caption(keep_with_next=True, bookmark=bm_name)
        make_image(keep_with_next=False)
    else:
        make_image(keep_with_next=True, bookmark=bm_name)
        make_caption(keep_with_next=False)


def _index_paragraph(doc, indent_cm: float, text: str, bookmark: str,
                     *, bold: bool = False, uppercase: bool = False) -> None:
    """Emit one entry of the manually built Index / List of figures / List of
    tables: ``<text>........<page-ref>`` with a right-aligned tab stop at the
    body width and a dot leader between text and page number."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    # Right tab at ~15.5 cm (body width inside the Deusto margins) with a dot
    # leader: produces the "Chapter title ............. 12" effect.
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    label = text.upper() if uppercase else text
    run = p.add_run(label)
    if bold:
        run.bold = True
    p.add_run("\t")
    _insert_pageref_field(p, bookmark)


def _add_front_matter_subtitle(doc, label: str) -> None:
    """Add an Abstract / Index / Figure index / Table index heading using the
    template's ``Subtitle`` style when available."""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles["Subtitle"]
    except KeyError:
        try:
            p.style = doc.styles["Subttulo"]
        except KeyError:
            pass
    p.add_run(label)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)


def _add_body_paragraph(doc, text: str) -> None:
    p = doc.add_paragraph()
    try:
        p.style = doc.styles["Normal"]
    except KeyError:
        pass
    add_runs(p, text)


def _add_page_break(doc) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def extract_body_elements(doc) -> list:
    """Detach every direct child of the body (paragraphs and tables) except
    the final ``w:sectPr``, and return them in document order. The body is
    left containing only the ``w:sectPr`` so subsequent ``add_*`` calls
    produce a fresh sequence that we will later place before these elements.
    """
    body = doc.element.body
    sect = None
    extracted: list = []
    for ch in list(body.iterchildren()):
        if ch.tag == qn("w:sectPr"):
            sect = ch
            continue
        body.remove(ch)
        extracted.append(ch)
    # Re-attach the sectPr at the end if it was there.
    if sect is not None:
        if sect not in list(body.iterchildren()):
            body.append(sect)
    return extracted


def build_front_matter(doc, cfg: dict) -> None:
    """Append the front matter (Abstract, Keywords, Index, LOF, LOT) to the
    currently empty body. The caller has already extracted the body elements
    and will reattach them after this call returns.
    """
    # --- Abstract -----------------------------------------------------------
    _add_front_matter_subtitle(doc, "Abstract")
    for para in cfg["abstract_paragraphs"]:
        _add_body_paragraph(doc, para)

    # --- Keywords -----------------------------------------------------------
    _add_front_matter_subtitle(doc, "Keywords")
    _add_body_paragraph(doc, cfg["descriptors"])

    _add_page_break(doc)

    # --- Index --------------------------------------------------------------
    _add_front_matter_subtitle(doc, "Index")
    for level, title, bm in TOC_ENTRIES:
        if level == 1:
            _index_paragraph(doc, 0.0, title, bm, bold=True, uppercase=True)
        elif level == 2:
            _index_paragraph(doc, 0.6, title, bm)
        else:
            _index_paragraph(doc, 1.2, title, bm)

    _add_page_break(doc)

    # --- List of figures ----------------------------------------------------
    if LOF_ENTRIES:
        _add_front_matter_subtitle(doc, "Figure index")
        for caption, bm in LOF_ENTRIES:
            _index_paragraph(doc, 0.0, caption, bm)
        _add_page_break(doc)

    # --- List of tables -----------------------------------------------------
    if LOT_ENTRIES:
        _add_front_matter_subtitle(doc, "Table index")
        for label, bm in LOT_ENTRIES:
            _index_paragraph(doc, 0.0, label, bm)
        _add_page_break(doc)


def append_body_elements(doc, elements: list) -> None:
    """Reattach the previously detached body elements after the front matter,
    keeping the final ``w:sectPr`` as the last child of the body."""
    body = doc.element.body
    sect = None
    for ch in list(body.iterchildren()):
        if ch.tag == qn("w:sectPr"):
            sect = ch
            body.remove(ch)
    for el in elements:
        body.append(el)
    if sect is not None:
        body.append(sect)


def prepend_front_matter(doc, elements: list) -> None:
    """Move *elements* (the appended front-matter blocks) to the beginning of
    the body, in the same order they were produced."""
    body = doc.element.body
    anchor = None
    for ch in body.iterchildren():
        if ch.tag == qn("w:sectPr"):
            continue
        anchor = ch
        break
    for el in elements:
        body.remove(el)
    if anchor is None:
        for el in elements:
            body.append(el)
        return
    for el in elements:
        anchor.addprevious(el)


def render_blocks(doc, blocks: list[dict], md_root: Path) -> None:
    for block in blocks:
        kind = block["kind"]
        if kind == "h1":
            add_paragraph(doc, block["text"], STYLE_H1)
        elif kind == "h2":
            add_paragraph(doc, block["text"], STYLE_H2)
        elif kind == "h3":
            add_paragraph(doc, block["text"], STYLE_H3)
        elif kind == "para":
            add_paragraph(doc, block["text"], STYLE_BODY)
        elif kind == "list":
            for item in block["items"]:
                add_paragraph(doc, item, STYLE_LIST)
        elif kind == "table":
            add_table(doc, block["header"], block["rows"])
        elif kind == "figure":
            add_figure(doc, block["caption"], block["path"], md_root)


def collect_blocks_for_chapter(entry, md_root: Path) -> list[dict]:
    """Return the merged block list for one chapter entry.

    A plain string is a single file. A tuple ``(folder, [files...])`` is a
    chapter split across multiple sub-files (Cap. 7 in this project). The
    first sub-file is expected to declare the ``# X. Title`` heading.
    """
    if isinstance(entry, str):
        text = (md_root / entry).read_text(encoding="utf-8")
        return parse_md(text)
    folder, files = entry
    merged: list[dict] = []
    for fname in files:
        text = (md_root / folder / fname).read_text(encoding="utf-8")
        merged.extend(parse_md(text))
    return merged


# ---------- main ---------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--grade", choices=["ing", "cdia"], required=True)
    parser.add_argument("--pdf", action="store_true",
                        help="After saving the .docx, open it in Word, refresh "
                             "fields and export to PDF in the same folder.")
    args = parser.parse_args()

    cfg = CONFIG[args.grade]
    docx_path: Path = cfg["docx_path"]
    md_root: Path = cfg["md_root"]

    if not docx_path.exists():
        raise SystemExit(f"docx not found: {docx_path}")

    print(f"Loading template: {docx_path}")
    doc = Document(str(docx_path))

    print("Promoting template sectPr to body level if needed...")
    ensure_final_sectpr(doc)

    print("Wiping body completely (covers + duplicated placeholders gone)...")
    wipe_body_completely(doc)

    # Reset all counters and metadata lists for this run.
    CAPTION_POSITION[0] = cfg.get("caption_position", "below")
    FIGURE_COUNTER[0] = 0
    TABLE_COUNTER[0] = 0
    H1_COUNTER[0] = 0
    BOOKMARK_COUNTER[0] = 1000
    TOC_ENTRIES.clear()
    LOF_ENTRIES.clear()
    LOT_ENTRIES.clear()

    for entry in cfg["chapter_files"]:
        label = entry if isinstance(entry, str) else entry[0]
        print(f"  rendering {label}")
        blocks = collect_blocks_for_chapter(entry, md_root)
        render_blocks(doc, blocks, md_root)

    print(f"  collected {len(TOC_ENTRIES)} index entries, "
          f"{len(LOF_ENTRIES)} figures, {len(LOT_ENTRIES)} tables")

    print("Extracting body elements to make room for the front matter...")
    body_elements = extract_body_elements(doc)
    print(f"  extracted {len(body_elements)} body elements")

    print("Building front matter (Abstract, Keywords, Index, LOF, LOT)...")
    build_front_matter(doc, cfg)

    print("Reattaching body elements after the front matter...")
    append_body_elements(doc, body_elements)

    # Pagination (covers without number, lower-roman frontmatter, decimal body)
    # is intentionally NOT configured here. Manipulating multiple sectPr from
    # python-docx breaks ``add_table`` when invoked later in the same run. The
    # operator configures pagination in Word in three clicks:
    #   1. Double-click the cover footer, untick "Link to Previous", clear it.
    #   2. Insert > Section break > Next page between LOT and chapter 1.
    #   3. Page Number > Format on each section (roman for frontmatter,
    #      decimal restarting at 1 for the body).

    print(f"Saving: {docx_path}")
    doc.save(str(docx_path))

    print("Patching settings.xml (updateFields=true)...")
    enable_update_fields(docx_path)

    print("Finalising indices through Word automation (TOC, LOF, LOT)...")
    try:
        finalize_docx_indices(docx_path)
        print("Indices populated; saved in place.")
    except Exception as exc:
        print(f"  warning: could not finalise indices automatically ({exc!s}).")
        print("  open the docx in Word and press Ctrl+A then F9 to update fields.")

    if args.pdf:
        print("Exporting to PDF through Word automation...")
        pdf = export_to_pdf(docx_path)
        print(f"PDF saved: {pdf}")

    print("Done.")


if __name__ == "__main__":
    main()
